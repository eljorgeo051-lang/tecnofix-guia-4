import os
import datetime
import site
from flask import Flask
from flask_sqlalchemy import SQLAlchemy
import customtkinter as ctk
from tkinter import messagebox, filedialog
from docx import Document
from docx.shared import Pt

# =============================================================================
# 1. BASE DE DATOS (PERSISTENCIA)
# =============================================================================
app = Flask(__name__)
db_path = os.path.join(os.path.abspath(os.path.dirname(__file__)), 'tecnofix_final.db')
app.config['SQLALCHEMY_DATABASE_URI'] = f'sqlite:///{db_path}'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)

class ClienteDB(db.Model):
    id = db.Column(db.String(20), primary_key=True)
    nombre = db.Column(db.String(100), nullable=False)

class OrdenDB(db.Model):
    id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    cliente_id = db.Column(db.String(20), db.ForeignKey('cliente_db.id'))
    tipo_equipo = db.Column(db.String(50))
    servicio = db.Column(db.String(100))
    total = db.Column(db.Float)
    fecha = db.Column(db.String(20))

with app.app_context():
    db.create_all()

# =============================================================================
# 2. CATÁLOGO DE SERVICIOS
# =============================================================================
CATALOGO_PRECIOS = {
    "📱 Celular": {
        "Diagnóstico técnico": 15000,
        "Cambio de pantalla (Económica)": 120000,
        "Cambio de pantalla (Original)": 250000,
        "Cambio de batería": 75000,
        "Reparación de software": 50000,
        "Cambio de pin de carga": 85000,
        "Mantenimiento preventivo": 30000,
        "Upgrade de almacenamiento": 140000
    },
    "💻 Computador": {
        "Diagnóstico técnico": 25000,
        "Formateo básico": 70000,
        "Instalación Windows + Programas": 80000,
        "Limpieza + Pasta Térmica": 65000,
        "Reparación tarjeta madre": 185000,
        "Aumento de memoria RAM": 70000,
        "Instalación de SSD": 80000,
        "Eliminación de virus/malware": 45000
    }
}

# =============================================================================
# 3. INTERFAZ GRÁFICA PROFESIONAL
# =============================================================================
class TecnoFixApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("TecnoFix - Gestión UNAD")
        self.geometry("600x820")
        ctk.set_appearance_mode("dark")

        # Título principal
        self.lbl_main = ctk.CTkLabel(self, text="🛠️ REGISTRO DE SERVICIOS TECNOFIX", font=("Roboto", 22, "bold"))
        self.lbl_main.pack(pady=(20, 5))

        # --- SECCIÓN DE DESARROLLADORES (ACTUALIZADA) ---
        self.frame_devs = ctk.CTkFrame(self, fg_color="transparent")
        self.frame_devs.pack(pady=(0, 15))
        
        # Jorge Gómez de primero con nombres completos
        autores = "Jorge Eliecer Gómez | Andres Felipe Rojas | Nicol Becerra"
        self.lbl_devs = ctk.CTkLabel(self.frame_devs, text=f"Desarrollado por: {autores}", 
                                     font=("Arial", 12, "italic"), text_color="#A9A9A9")
        self.lbl_devs.pack()

        # Frame de Datos del Cliente
        self.frame = ctk.CTkFrame(self)
        self.frame.pack(pady=10, padx=40, fill="both", expand=True)

        self.ent_id = ctk.CTkEntry(self.frame, placeholder_text="Cédula del Cliente")
        self.ent_id.pack(pady=10, padx=20, fill="x")
        
        self.ent_nom = ctk.CTkEntry(self.frame, placeholder_text="Nombre Completo")
        self.ent_nom.pack(pady=10, padx=20, fill="x")

        # Selección de Servicio
        self.combo_eq = ctk.CTkComboBox(self.frame, values=list(CATALOGO_PRECIOS.keys()), command=self.actualizar_servicios)
        self.combo_eq.pack(pady=15)

        self.combo_srv = ctk.CTkComboBox(self.frame, values=[], width=350, command=self.mostrar_precio)
        self.combo_srv.pack(pady=10)

        # Precio visual automático[cite: 1]
        self.lbl_precio_display = ctk.CTkLabel(self.frame, text="$ 0", font=("Arial", 36, "bold"), text_color="#2fa572")
        self.lbl_precio_display.pack(pady=20)

        # Botón de Acción
        self.btn_guardar = ctk.CTkButton(self, text="GUARDAR Y GENERAR FACTURA WORD", 
                                         height=50, font=("Arial", 14, "bold"), 
                                         fg_color="#1f538d", hover_color="#14375e",
                                         command=self.finalizar_registro)
        self.btn_guardar.pack(pady=20, padx=40, fill="x")

        # Créditos Institucionales
        self.lbl_unad = ctk.CTkLabel(self, text="Programación de Computadores 213023 - UNAD", font=("Arial", 10))
        self.lbl_unad.pack(pady=(0, 15))

        self.actualizar_servicios(self.combo_eq.get())

    def actualizar_servicios(self, eq):
        servicios = list(CATALOGO_PRECIOS[eq].keys())
        self.combo_srv.configure(values=servicios)
        self.combo_srv.set(servicios[0])
        self.mostrar_precio(servicios[0])

    def mostrar_precio(self, srv):
        eq = self.combo_eq.get()
        precio = CATALOGO_PRECIOS[eq][srv]
        self.lbl_precio_display.configure(text=f"$ {precio:,.0f} COP")

    def finalizar_registro(self):
        cid = self.ent_id.get()
        nom = self.ent_nom.get()
        eq = self.combo_eq.get()
        srv = self.combo_srv.get()
        precio = CATALOGO_PRECIOS[eq][srv]
        fecha = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")

        if not cid or not nom:
            messagebox.showwarning("Atención", "Por favor ingrese los datos del cliente.")
            return

        # Selección de destino del archivo[cite: 1]
        ruta_archivo = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Documento de Word", "*.docx")],
            initialfile=f"Factura_{nom.replace(' ', '_')}"
        )

        if not ruta_archivo:
            return

        try:
            with app.app_context():
                # Registro en BD[cite: 1]
                if not ClienteDB.query.get(cid):
                    db.session.add(ClienteDB(id=cid, nombre=nom))
                
                nueva = OrdenDB(cliente_id=cid, tipo_equipo=eq, servicio=srv, total=precio, fecha=fecha)
                db.session.add(nueva)
                db.session.commit()
                
                # Construcción del Documento Word[cite: 1]
                doc = Document()
                doc.add_heading('TECNOFIX - COMPROBANTE DE SERVICIO', 0)
                
                doc.add_paragraph(f"Orden de Trabajo N°: {nueva.id}")
                doc.add_paragraph(f"Fecha: {fecha}")
                doc.add_paragraph("-" * 50)
                
                doc.add_heading('Información del Cliente', level=1)
                doc.add_paragraph(f"Nombre: {nom}\nDocumento: {cid}")
                
                doc.add_heading('Detalles Técnicos', level=1)
                doc.add_paragraph(f"Categoría: {eq}\nServicio Prestado: {srv}")
                
                # Precio total destacado[cite: 1]
                p_total = doc.add_paragraph()
                p_total.add_run(f"TOTAL A PAGAR: ${precio:,.0f} COP").bold = True
                
                # Créditos de autores en el documento[cite: 1]
                doc.add_paragraph("\n" + "_"*40)
                doc.add_paragraph("Sistema desarrollado por: Jorge Eliecer Gómez Rojas")
                doc.add_paragraph("Proyecto Académico UNAD - Ingeniería de Sistemas.")
                
                doc.save(ruta_archivo)

            messagebox.showinfo("Éxito", f"La factura se ha generado correctamente.")
            self.ent_id.delete(0, 'end')
            self.ent_nom.delete(0, 'end')

        except Exception as e:
            messagebox.showerror("Error", f"No se pudo completar la operación: {e}")

if __name__ == "__main__":
    TecnoFixApp().mainloop()